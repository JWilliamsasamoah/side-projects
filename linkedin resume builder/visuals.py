import streamlit as st
from dotenv import load_dotenv
import os
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import time
import json
from groq import Groq
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, RGBColor
import re 
# Load environment variables
load_dotenv()



def create_resume(generated_resume, filename="Final_Resume.docx"):
    doc = Document()

    # Helper function for styling
    def set_run_style(run, font_size=12, bold=False, color=None):
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_section_title(doc, title):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(title)
        set_run_style(run, font_size=12, bold=True, color=(0, 112, 192))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_bullet_paragraph(doc, text):
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(text)
        set_run_style(run, font_size=10)

    # Split the AI-provided content into sections
    lines = generated_resume.strip().split("\n")
    sections = {"NAME": "", "CONTACT": "", "OBJECTIVE": "", "EXPERIENCE": "", "EDUCATION": "", "SKILLS": "", "PROJECTS": ""}
    current_section = None

    for line in lines:
        line = line.strip()
        if line.upper() in sections:  # Check for section headers
            current_section = line.upper()
        elif current_section:
            sections[current_section] += line + "\n"
        elif current_section is None and len(line.split()) <= 3:  # Likely a name
            sections["NAME"] += line.strip()

    # Add Name and Contact
    doc.add_paragraph(sections["NAME"], style="Heading 1")
    doc.add_paragraph(sections["CONTACT"])

    # Add Objective Section
    if sections["OBJECTIVE"].strip():
        add_section_title(doc, "OBJECTIVE")
        doc.add_paragraph(sections["OBJECTIVE"].strip())

    # Add Experience Section
    if sections["EXPERIENCE"].strip():
        add_section_title(doc, "EXPERIENCE")
        for exp in sections["EXPERIENCE"].strip().split("\n\n"):
            exp_details = exp.split("\n")
            if len(exp_details) >= 2:
                doc.add_paragraph(f"{exp_details[0]} — {exp_details[1]}")
                for detail in exp_details[2:]:
                    add_bullet_paragraph(doc, detail)

    # Add Education Section
    if sections["EDUCATION"].strip():
        add_section_title(doc, "EDUCATION")
        for edu in sections["EDUCATION"].strip().split("\n\n"):
            edu_details = edu.split("\n")
            if len(edu_details) >= 2:
                doc.add_paragraph(f"{edu_details[0]} — {edu_details[1]}")
                for detail in edu_details[2:]:
                    add_bullet_paragraph(doc, detail)

    # Add Skills Section
    if sections["SKILLS"].strip():
        add_section_title(doc, "SKILLS")
        doc.add_paragraph(sections["SKILLS"].strip())

    # Add Projects Section
    if sections["PROJECTS"].strip():
        add_section_title(doc, "PROJECTS")
        for project in sections["PROJECTS"].strip().split("\n\n"):
            doc.add_paragraph(project.strip())

    doc.save(filename)
    return filename



# Function to scrape LinkedIn sections
def scrape_section(section_name, xpath_pattern, driver):
    try:
        section = driver.find_element(By.XPATH, f"//section[.//span[text()='{section_name}']]")
        items = section.find_elements(By.XPATH, xpath_pattern)
        data = []
        for item in items:
            item_data = {}
            try:
                item_data['title'] = item.find_element(By.XPATH, ".//div[contains(@class, 't-bold')]/span[1]").text
            except:
                item_data['title'] = ""
            try:
                item_data['subtitle'] = item.find_element(By.XPATH, ".//span[contains(@class, 't-normal')]").text
            except:
                item_data['subtitle'] = ""
            try:
                item_data['dates'] = item.find_element(By.XPATH, ".//span[contains(@class, 't-black--light')]").text
            except:
                item_data['dates'] = ""
            try:
                item_data['description'] = item.find_element(By.XPATH, ".//div[contains(@class, 'inline-show-more-text')]").text
            except:
                item_data['description'] = ""
            data.append(item_data)
        return data
    except:
        return []
# Function to scrape LinkedIn Skills
def scrape_skills_page(driver, profile_url):
    skills_data = []
    try:
        skills_url = profile_url.rstrip("/") + "/details/skills/"
        driver.get(skills_url)
        time.sleep(3)

        skills_items = driver.find_elements(By.XPATH, "//li[contains(@class, 'pvs-list__paged-list-item')]")

        for item in skills_items:
            skill_info = {}
            try:
                skill_name = item.find_element(By.XPATH, ".//span[@aria-hidden='true']").text
                skill_info["name"] = skill_name.strip() if skill_name else ""
            except:
                continue

            try:
                skill_detail = item.find_element(By.XPATH, ".//span[contains(@class, 't-14') and not(@aria-hidden='true')]").text
                skill_info["detail"] = skill_detail.strip() if skill_detail else ""
            except:
                skill_info["detail"] = ""

            if skill_info not in skills_data:
                skills_data.append(skill_info)

    except:
        pass
    return skills_data

# Streamlit App
st.markdown(
    """
    <style>
        body {
            background-color: #E3F2FD;
        }
        .stTextInput>div>div>div>input, .stTextArea>div>div>textarea {
            color: white;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #0D47A1;
        }
        .stButton>button {
            background-color: #0D47A1;
            color: white;
        }
    </style>
    """,
    unsafe_allow_html=True,
)
st.title("LinkedIn Resume Helper")
st.markdown("Extract LinkedIn profile data and generate a resume")
progress = st.progress(0)
# Input fields
email = st.text_input("LinkedIn Email", placeholder="Enter your LinkedIn email")
password = st.text_input("LinkedIn Password", placeholder="Enter your LinkedIn password", type="password")
profile_url = st.text_input("LinkedIn Profile URL", placeholder="Enter the LinkedIn profile URL")
generate = st.button("Generate your Resume")

if generate:
    # Validation
    if not email or not password or not profile_url:
        st.error("Please provide all required fields!")
    else:
        st.info("Starting LinkedIn scraping...")

        # Initialize Selenium WebDriver
        options = webdriver.ChromeOptions()
        options.add_argument("--log-level=3")
        options.add_experimental_option("excludeSwitches", ["enable-logging"])
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

        try:
            # Log in to LinkedIn
            driver.get("https://www.linkedin.com/login")
            progress.progress(20)
            driver.find_element(By.ID, "username").send_keys(email)
            driver.find_element(By.ID, "password").send_keys(password)
            driver.find_element(By.ID, "password").send_keys(Keys.RETURN)
            time.sleep(10)
            progress.progress(40)

            # Navigate to profile
            driver.get(profile_url)
            time.sleep(3)
            progress.progress(60)

            # Initialize variables to store data
            basic_info = {}
            experience_data = []
            education_data = []
            licenses_data = []
            projects_data = []
            skills_data = []

            # Extract profile information
            try:
                # Extract basic information
                print("\nBasic Information:")
                try:
                    basic_info["name"] = driver.find_element(By.CSS_SELECTOR, "h1.text-heading-xlarge").text
                    basic_info["headline"] = driver.find_element(By.CSS_SELECTOR, "div.text-body-medium.break-words").text
                    basic_info["location"] = driver.find_element(By.CSS_SELECTOR, "span.text-body-small.inline.t-black--light.break-words").text
                    print(basic_info)
                except:
                    print("Basic information not available.")
                progress.progress(70)

                # Scroll down to load more sections
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                time.sleep(2)

                # Extract Experience
                experience_data = scrape_section("Experience", ".//li[contains(@class, 'artdeco-list__item')]", driver)
                if experience_data:
                    print("\nExperience:")
                    for exp in experience_data:
                        print(exp)

                # Extract Education
                education_data = scrape_section("Education", ".//li[contains(@class, 'artdeco-list__item')]", driver)
                if education_data:
                    print("\nEducation:")
                    for edu in education_data:
                        print(edu)

                # Extract Licenses & Certifications
                licenses_data = scrape_section("Licenses & certifications", ".//li[contains(@class, 'artdeco-list__item')]", driver)
                if licenses_data:
                    print("\nLicenses & Certifications:")
                    for license in licenses_data:
                        print(license)

                # Extract Projects
                projects_data = scrape_section("Projects", ".//li[contains(@class, 'artdeco-list__item')]", driver)
                if projects_data:
                    print("\nProjects:")
                    for project in projects_data:
                        print(project)

                # Extract Skills
                skills_data = scrape_skills_page(driver, profile_url)
                if skills_data:
                    print("\nSkills:")
                    for skill in skills_data:
                        print(f"Skill: {skill['name']}, Detail: {skill['detail']}")

            except Exception as e:
                print(f"Error extracting information: {e}")
                
            # Close the browser
            driver.quit()
            progress.progress(90)
            # Combine scraped data
            alldata = {
                "basic_info": basic_info,
                "experience": experience_data,
                "education": education_data,
                "licenses_and_certifications": licenses_data,
                "projects": projects_data,
                "skills": skills_data,
            }

            # Display scraped data
            st.success("Scraping complete!")
            st.json(alldata)
            progress.progress(100)

            # Generate resume using Groq API
            st.info("Generating resume...")
            client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
            try:
                chat_completion = client.chat.completions.create(
                    messages=[
                        {
                            "role": "user",
                            "content": """Please create a professional resume using the following information. Format it neatly and structure it into sections: "Personal Information," "Experience," "Education," "Skills," and "Projects." Follow these formatting rules:

1. Do NOT use symbols like asterisks (`**`) for the title or any where in the texts , underscores (`_`), or any special characters for bolding or emphasis. Instead, rely on capitalization and proper indentation to structure the text.

2. Section Titles:
   - Write section titles (e.g., "EXPERIENCE," "EDUCATION," "SKILLS") in capital letters.
   - Leave one blank line before each section title for clarity.

3. Personal Information:
   - Place the individual's name prominently at the top, followed by their professional title, location, phone number, and email address.

4. OBjective :
    Make an Objective for the resume using the the info
4. Experience:
   - List each job in reverse chronological order (most recent first).
   - For each job, include:
     - Job title and company name on the same line, separated by a dash (e.g., "FIRST Robotics Coach & Programmer - Education Bank").
     - Dates of employment on the right-hand side.
     - A bulleted list of responsibilities and achievements for each role. Use a hyphen (`-`) to start each bullet.

5. Education:
   - List degrees and institutions in reverse chronological order.
   - For each entry, include:
     - Degree and field of study, followed by the institution name (e.g., "BSc in Computer Science - Example University").
     - Dates of attendance on the right-hand side.

6. Skills:
   - Use a comma-separated list for skills (e.g., "Python, JavaScript, Machine Learning").

7. Projects:
   - List relevant projects with their name, timeline, and a description of accomplishments. Use the same formatting as in the "Experience" section.

8. General Layout:
   - Maintain consistent alignment and spacing throughout the resume.
   - Ensure the text is well-structured, easy to read, and professional in appearance. as well use this info to make an objecti

Here is the data to be used for the resume:
\n""" + json.dumps(alldata, indent=4),
                        }
                    ],
                    model="llama3-8b-8192",
                )
                st.success("Resume generated!")
                st.text_area("Generated Resume", chat_completion.choices[0].message.content, height=500)
                generated_resume = chat_completion.choices[0].message.content
                generated_resume = generated_resume.replace("**", "").strip()


               

                filename = create_resume(generated_resume)
                # Offer download link in Streamlit
                with open(filename, "rb") as file:
                    st.download_button(
                        label="Download Resume",
                        data=file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            except Exception as e:
                st.error(f"Error: {e}")
                
                print(generated_resume)
                
        finally:
            driver.quit()